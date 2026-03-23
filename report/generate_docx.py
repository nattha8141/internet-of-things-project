"""
Generate the final report .docx from the template.
Uses python-docx to write content into the Imperial template.
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os

TEMPLATE_PATH = "coursework/template.docx"
OUTPUT_PATH = "report/final_report.docx"
IMAGE_DIR = "images/"

def add_heading_styled(doc, text, level=1):
    doc.add_heading(text, level=level)

def add_para(doc, text, bold=False, italic=False, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    if bold:
        run.bold = True
    if italic:
        run.italic = True
    return p

def add_image_with_caption(doc, image_path, caption, width=5.5):
    if os.path.exists(image_path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(image_path, width=Inches(width))
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cap.add_run(caption)
        run.italic = True
        run.font.size = Pt(9)

def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(9)
    # Data rows
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)
    return table

def main():
    # Create from scratch (template has incompatible styles)
    doc = Document()
    print("Creating report document from scratch")

    # Clear existing content from template (keep styles)
    # We'll just append after whatever is in the template

    # ========== TITLE / HEADER ==========
    doc.add_paragraph()  # spacer
    title = doc.add_heading('The Effect of Colored Light on Bean Sprout Growth', level=0)
    add_para(doc, 'CID: 06043088 | ELEC70126 — Internet of Things and Applications', bold=True, size=12)
    add_para(doc, '')

    # ========== PART 1: SENSING ==========
    doc.add_heading('Part 1: Sensing', level=1)

    # --- 1.1 Introduction ---
    doc.add_heading('1. Introduction and Objectives', level=2)
    add_para(doc,
        'Mung bean sprouts (Vigna radiata) are a widely cultivated crop known for their rapid germination '
        'and vigorous growth in dark conditions. This experiment investigates how exposure to different '
        'wavelengths of visible light — specifically blue (~450 nm) and green (~520 nm) — affects the rate '
        'of sprout growth compared to a dark control group.')
    add_para(doc, '')
    add_para(doc, 'Objectives:', bold=True)
    objectives = [
        'Quantify the growth deceleration (or acceleration) of mung bean sprouts under blue and green light relative to a dark control.',
        'Design and deploy an end-to-end IoT sensing system using ESP32, photoresistors, and environmental sensors.',
        'Collect continuous time-series data over a week-long experiment with automated cloud-based storage.',
        'Identify any unexpected plant behaviors (e.g., phototropic movement) observable through the sensor data.'
    ]
    for obj in objectives:
        doc.add_paragraph(obj, style='List Bullet')

    add_para(doc,
        'Hypothesis: Blue light, known to stimulate phototropin receptors and influence photomorphogenesis, '
        'is expected to induce a stronger growth response than green light, while the dark control should '
        'exhibit uninhibited etiolated growth.', italic=True)

    # --- 1.2 Data Sources ---
    doc.add_heading('2. Data Sources and Sensing Set-Up', level=2)
    doc.add_heading('2.1 Physical Setup', level=3)
    add_para(doc,
        'The experiment uses a divided cardboard enclosure partitioned into three isolated chambers:')

    add_table(doc,
        ['Chamber', 'Condition', 'Light Source'],
        [
            ['Left', 'Blue light (continuous)', 'Blue LED (GPIO 47, PWM 255)'],
            ['Centre', 'Control / Dark', 'No light (1s red pulse every 15 min for measurement)'],
            ['Right', 'Green light (continuous)', 'Green LED (GPIO 48, PWM 50)'],
        ])

    add_para(doc,
        'Each chamber contains a small plastic container with mung beans placed on wet cotton wool '
        'as the growth medium. Cardboard dividers prevent light contamination between chambers.')

    add_image_with_caption(doc, IMAGE_DIR + 'day_1.jpeg',
        'Figure 1: Experiment setup on Day 1 — three chambers with mung beans on cotton wool.')

    add_image_with_caption(doc, IMAGE_DIR + 'medium_and_environment.jpeg',
        'Figure 2: Close-up of growth medium and LED/sensor arrangement.')

    doc.add_heading('2.2 Sensor Configuration', level=3)
    add_para(doc,
        'The system employs two types of time-series data sources of different nature:')
    add_para(doc, '1. Photoresistors (LDRs) — 3 units', bold=True)
    add_para(doc,
        'Connected to ESP32 ADC pins (GPIO 1, 2, 4). Each LDR is paired with a red measurement LED. '
        'As sprouts grow and obstruct the optical path, the ADC reading increases. The red LED is activated '
        'for 5 seconds every 15 minutes, with ADC sampling 2 seconds into the pulse. Resolution: 12-bit (0–4095).')
    add_para(doc, '2. DHT11 Temperature and Humidity Sensor — 1 unit', bold=True)
    add_para(doc,
        'Connected to GPIO 5, monitoring ambient temperature (°C) and relative humidity (%). '
        'Provides environmental context and enables future hydration-based actuation.')

    doc.add_heading('2.3 Design Choices', level=3)
    add_table(doc,
        ['Decision', 'Choice', 'Justification'],
        [
            ['Measurement interval', '15 minutes', 'Balances resolution (96 samples/day) with minimal light exposure'],
            ['Measurement LED', 'Red, 5s pulse', 'Red has minimal effect on growth; short pulse limits exposure'],
            ['ADC read timing', '2s into pulse', 'Allows photoresistor to stabilise'],
            ['Environmental sensor', 'DHT11', 'Adequate for trend detection (±2°C, ±5% RH); low cost'],
            ['Microcontroller', 'ESP32', 'Integrated Wi-Fi, sufficient ADC channels, low cost (~£5)'],
        ])

    # --- 1.3 Data Collection ---
    doc.add_heading('3. Data Collection and Storage/Communications Process', level=2)
    add_para(doc, 'Data Pipeline Architecture:', bold=True)
    add_para(doc, 'ESP32 Sensors → Wi-Fi (HTTP GET) → Google Apps Script → Google Sheets → CSV Export')
    add_para(doc, '')
    add_para(doc,
        'Every 15 minutes, the ESP32 activates the red measurement LED, waits 2 seconds, reads all 3 '
        'photoresistors and the DHT11, then sends the data via HTTP GET to a Google Apps Script endpoint. '
        'The script appends each data point as a new row in Google Sheets, providing cloud-based persistent '
        'storage with real-time accessibility.')

    add_para(doc, 'Data Cleaning:', bold=True)
    add_para(doc,
        'Two types of anomalies were addressed: (1) Physical disturbance artefacts — sudden simultaneous drops '
        'in all three channels caused by opening the enclosure for watering. These were detected when |Δ| > 200 '
        'occurred in all three channels simultaneously, and replaced via linear interpolation. '
        '(2) Sensor saturation — the blue channel reached ADC max (4095) ~3 days in.')

    add_table(doc,
        ['Parameter', 'Value'],
        [
            ['Sampling interval', '15 minutes'],
            ['Total data points', '558'],
            ['Duration', '5 days 20 hours (5–10 Mar 2026)'],
            ['Channels', '5 (Green, Blue, Control ADC + Temp + Humidity)'],
            ['Data size', '~30 KB (CSV)'],
        ])

    # --- 1.4 Basic Characteristics ---
    doc.add_heading('4. Basic Characteristics of the End-to-End Setup', level=2)

    add_table(doc,
        ['Metric', 'Value'],
        [
            ['Sampling rate', '1 sample / 15 min (96/day)'],
            ['Throughput', '~4.8 KB/day'],
            ['Latency (sensor → cloud)', '2–5 seconds'],
            ['Power', '~150 mA (USB-powered)'],
            ['Uptime', '>99%'],
            ['ADC resolution', '12-bit (0–4095)'],
        ])

    add_para(doc, 'Trade-Off Analysis:', bold=True)
    add_table(doc,
        ['Trade-Off', 'Decision', 'Rationale'],
        [
            ['Local vs Cloud', 'Cloud (Google Sheets)', 'Simplifies firmware; remote monitoring; ~5 KB/day'],
            ['Sampling vs Exposure', '15-min interval', 'Higher rates increase red LED exposure to control'],
            ['Accuracy vs Cost', 'DHT11 over BME280', 'Trend detection sufficient; £1 vs £8'],
            ['Reliability vs Simplicity', 'No local buffer', 'Acceptable for stable Wi-Fi environment'],
            ['Total cost', '~£15', 'ESP32+sensors+LEDs+breadboard+RTC'],
        ])

    add_para(doc, 'Limitations:', bold=True)
    limitations = [
        'Blue channel ADC saturation prevents full growth quantification.',
        'Ambient green light interference: the continuously-on green LED can reach the photoresistor even when '
        'sprouts partially obstruct the red measurement path, underrepresenting green group growth.',
        'Phototropic bending in the control group confounds growth readings — sprouts move out of the sensor '
        'path, causing apparent decreases that reflect movement, not reduced growth.',
        'Single environmental sensor; micro-climate may vary between chambers.',
        'No local buffering; Wi-Fi dropouts cause data loss.',
        'Breadboard connections are fragile; disturbances during watering cause anomalies.',
    ]
    for lim in limitations:
        doc.add_paragraph(lim, style='List Bullet')

    # ========== PART 2: IoT ==========
    doc.add_page_break()
    doc.add_heading('Part 2: Internet of Things', level=1)

    # --- 2.1 Platform ---
    doc.add_heading('1. Data Interaction/Visualisation/Actuation Platform', level=2)
    add_para(doc,
        'An interactive web dashboard was built using Streamlit (Python) with Plotly for visualisation, '
        'providing five view modes:')
    views = [
        'Dashboard: Overview of all growth curves with raw/cleaned toggle and baseline normalisation.',
        'Growth Explorer: Configurable growth rate with adjustable window and smoothing; daily bar charts.',
        'Movement Analysis: Detrending tool for plant oscillation with FFT frequency spectrum.',
        'Environmental: Temperature/humidity time series and scatter plots vs growth rate.',
        'Compare & Stats: Welch\'s t-test, box plots, and cross-correlation heatmaps.',
    ]
    for v in views:
        doc.add_paragraph(v, style='List Bullet')

    add_para(doc, 'To run: streamlit run app/bean_sprout_dashboard.py', italic=True)

    # --- 2.2 Analytics ---
    doc.add_heading('2. Data Analytics, Inferences and Insights', level=2)

    add_image_with_caption(doc, IMAGE_DIR + 'growth_curves.png',
        'Figure 3: Growth curves showing photoresistor readings over time. Blue saturates at 4095.')

    add_para(doc, 'Growth Comparison:', bold=True)
    add_table(doc,
        ['Group', 'Baseline', 'Final', 'Total Change', 'Notes'],
        [
            ['Green', '3479', '~3528', '+49', 'Moderate growth, underrepresented by sensor'],
            ['Blue', '3491', '4095', '+604*', 'Saturated — true growth exceeds this'],
            ['Control', '3595', '~3045', '-550', 'Growing normally, but bending away from sensor'],
        ])

    add_para(doc, 'Key Findings:', bold=True)
    findings = [
        'Blue light dramatically accelerated growth, saturating the sensor in ~3 days. This is consistent '
        'with blue light activating phototropin and cryptochrome receptors. Manual observation confirmed '
        'tallest sprouts with green, chlorophyll-rich leaves.',
        'Green light produced decent but sensor-underrepresented growth (+49 ADC). The ambient green photons '
        'can leak through to the photoresistor even when sprouts obstruct the red measurement path, keeping '
        'readings artificially low. Visual inspection confirmed reasonable growth with slightly green leaves.',
        'The control group grew normally but bent sideways (phototropism toward blue light leaks), moving '
        'sprouts out of the sensor path and causing apparent ADC decreases. The ±500 ADC oscillations '
        'reflect the sprouts swaying as they search for light.',
    ]
    for f in findings:
        doc.add_paragraph(f, style='List Bullet')

    add_para(doc, 'Leaf Coloration (Manual Observation):', bold=True)
    add_table(doc,
        ['Group', 'Leaf Appearance', 'Interpretation'],
        [
            ['Blue', 'Green, healthy leaves', 'Blue light activates chlorophyll biosynthesis'],
            ['Green', 'Slight hint of green', 'Partial photomorphogenic response'],
            ['Control', 'Yellowish, etiolated', 'No light stimulus for chlorophyll production'],
        ])

    # Geometric Height Model
    add_para(doc, 'Geometric Height Estimation:', bold=True)
    add_para(doc,
        'Since no direct height sensor was available, we estimated plant height from ADC readings using '
        'the chamber geometry. The LED sits at ~11 cm height, the photosensor at ~4 cm. Using trigonometry, '
        'we calculate that full light-path blockage (ADC = 4095) corresponds to a plant height of ~5.2 cm '
        'above the container rim. The obstruction ratio maps linearly to estimated height.')

    add_table(doc,
        ['Group', 'Est. Final Height', 'Peak Height', 'Notes'],
        [
            ['Blue', '5.2 cm (saturated)', '>=5.2 cm', 'Reached sensor limit in ~3 days'],
            ['Green', '~0.4 cm', '~0.6 cm', 'Underestimated due to green light leak'],
            ['Control', '~0 cm (oscillating)', '~2.5 cm', 'Movement dominates reading'],
        ])

    add_image_with_caption(doc, IMAGE_DIR + 'geometric_model.png',
        'Figure 4: ADC-to-height calibration curve (left) and chamber side-view geometry (right).')

    add_image_with_caption(doc, IMAGE_DIR + 'estimated_height_growth.png',
        'Figure 5: Estimated plant height over time and growth rate in cm/hour from geometric model.')

    add_image_with_caption(doc, IMAGE_DIR + 'plant_movement_oscillation.png',
        'Figure 6: Detrended signals revealing plant oscillation, especially prominent in the control group.')

    add_image_with_caption(doc, IMAGE_DIR + 'daily_growth.png',
        'Figure 5: Daily growth comparison across the three groups.')

    add_para(doc, 'Statistical Significance:', bold=True)
    add_para(doc,
        "Welch's t-tests on hourly growth rates confirm statistically significant differences between "
        "all group pairs (p < 0.001), validating that observed differences are not due to random variation.")

    add_para(doc, 'Note on Cross-Correlation:', bold=True)
    add_para(doc,
        'Cross-correlation between isolated chambers has limited causal meaning — any observed correlation '
        'is driven by shared environmental factors (temperature, time of day), not direct interaction. '
        'The more meaningful analysis is each channel vs environmental factors and temporal auto-correlation '
        'within each channel to identify periodicity.')

    add_image_with_caption(doc, IMAGE_DIR + 'correlation_matrix.png',
        'Figure 6: Cross-correlation matrix (note: inter-chamber correlation reflects shared environment, not causation).')

    add_image_with_caption(doc, IMAGE_DIR + 'environmental_conditions.png',
        'Figure 7: Environmental conditions (temperature and humidity) during the experiment.')

    # --- 2.3 Discussion ---
    doc.add_heading('3. Discussions on Important Aspects', level=2)

    add_para(doc, 'Innovation and Creativity:', bold=True)
    innovations = [
        'Indirect growth measurement using photoresistors as a low-cost, continuous, non-contact growth proxy.',
        'Serendipitous discovery of plant movement patterns from always-on IoT monitoring — phenomena that '
        'periodic manual observation would miss.',
        'Multi-insight from a single sensor: growth trends (long-term) and movement dynamics (short-term).',
    ]
    for inn in innovations:
        doc.add_paragraph(inn, style='List Bullet')

    add_para(doc, 'Scalability:', bold=True)
    add_para(doc,
        'The system scales horizontally (more ADC channels), across devices (multiple ESP32s to one Sheet), '
        'and to production cloud services (Firebase, InfluxDB, AWS IoT Core) given the HTTP architecture.')

    add_para(doc, 'Complexity:', bold=True)
    add_para(doc,
        'The system integrates hardware (ESP32, 3×LDR, DHT11, RTC, LEDs), firmware (Arduino C++), '
        'cloud services (Google Apps Script + Sheets), analytics (Python/Pandas/SciPy in Jupyter), '
        'and a web dashboard (Streamlit/Plotly) — covering the full IoT pipeline.')

    add_image_with_caption(doc, IMAGE_DIR + 'day_3.jpeg',
        'Figure 8: Bean sprout growth on Day 3.')
    add_image_with_caption(doc, IMAGE_DIR + 'day_5.jpeg',
        'Figure 9: Bean sprout growth on Day 5 — significant growth visible across all chambers.')

    # --- 2.4 Future Work ---
    doc.add_heading('4. Avenues for Future Work and Potential Impact', level=2)

    add_para(doc, 'Technical Improvements:', bold=True)
    tech = [
        'Higher-resolution ADC or attenuated LEDs to prevent sensor saturation.',
        'SD card buffering for data resilience during Wi-Fi outages.',
        'Automated watering actuation triggered by humidity thresholds — completing the IoT feedback loop.',
        'Camera integration for time-lapse visual verification.',
    ]
    for t in tech:
        doc.add_paragraph(t, style='List Bullet')

    add_para(doc, 'Scientific Extensions:', bold=True)
    science = [
        'Longer experiments (2–4 weeks) covering full growth cycle.',
        'Additional wavelengths (red, UV, white) for complete action spectrum.',
        'Dedicated multi-angle photoresistors for 2D/3D plant movement tracking.',
        'Colour sensor (e.g., TCS34725) to quantify leaf chlorophyll content over time — linking light '
        'wavelength to photosynthetic development, not just stem elongation.',
    ]
    for s in science:
        doc.add_paragraph(s, style='List Bullet')

    add_para(doc, 'Potential Impact:', bold=True)
    impact = [
        'Urban/vertical farming: IoT light recipes to optimise indoor crop growth.',
        'Education: Low-cost STEM teaching tool for plant biology, electronics, and data science.',
        'Plant phenotyping: Continuous movement monitoring for agricultural research.',
    ]
    for im in impact:
        doc.add_paragraph(im, style='List Bullet')

    add_image_with_caption(doc, IMAGE_DIR + 'day_7.jpeg',
        'Figure 10: Bean sprout growth on Day 7 — dense growth with visible phototropic bending.')

    # --- References ---
    add_para(doc, '')
    add_para(doc, 'Code Repository: https://github.com/hajidnaufalatthousi/internet-of-things-project', bold=True, size=10)
    add_para(doc, 'Data: https://docs.google.com/spreadsheets/d/1Vkc24L-VDzpiR6GrKL9sg7BJ5h5271bmTEASLvmzD9M/', size=10)

    # Save
    os.makedirs(os.path.dirname(OUTPUT_PATH), exist_ok=True)
    doc.save(OUTPUT_PATH)
    print(f"Report saved to: {OUTPUT_PATH}")

if __name__ == "__main__":
    main()
